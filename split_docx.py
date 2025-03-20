#!/usr/bin/env python3
import sys
import os
import datetime
import logging
from docx import Document
from docx.oxml.ns import qn

# --------------------------------------------------
# Configuration
# --------------------------------------------------
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
ERROR_LOG_FILE = "split_error.log"

# --------------------------------------------------
# Helper Functions
# --------------------------------------------------
def record_error(error_type, location, original_text, error_message):
    """
    Append an error entry to the error log file.
    """
    timestamp = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    log_entry = f"[{timestamp}] {error_type} Error at {location}: {error_message}\n"
    if original_text:
        log_entry += f"   Original Text: {original_text}\n"
    try:
        with open(ERROR_LOG_FILE, "a", encoding="utf-8") as f:
            f.write(log_entry)
    except Exception as ex:
        logging.error(f"Failed to write to error log: {ex}")


def run_has_page_break(run):
    """
    Check if the run contains a page break by examining its XML.
    """
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    try:
        for br in run._element.findall('.//w:br', namespaces=nsmap):
            if br.get(qn('w:type')) == 'page':
                return True
    except Exception as e:
        logging.error(f"Error checking page break: {e}")
        record_error("PageBreak", "run_has_page_break", "", str(e))
    return False


def count_page_breaks(paragraph):
    """
    Count the number of page breaks present in a paragraph.
    """
    count = 0
    try:
        for run in paragraph.runs:
            if run_has_page_break(run):
                count += 1
    except Exception as e:
        logging.error(f"Error counting page breaks: {e}")
        record_error("Splitting", "Paragraph page break count", paragraph.text, str(e))
    return count


def split_docx(input_path, segment_page_count, output_dir):
    """
    Split the input DOCX into segments of approximately 'segment_page_count' pages each.
    Returns a list of file paths for the splitted segments.
    """
    try:
        doc = Document(input_path)
    except Exception as e:
        logging.error(f"Could not open input DOCX {input_path}: {e}")
        record_error("Splitting", "Opening document", "", str(e))
        return []

    if not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    splitted_files = []
    global_page = 1
    segment_start_page = 1
    segment_index = 1
    current_segment_doc = Document()

    for para_index, para in enumerate(doc.paragraphs):
        try:
            current_segment_doc.add_paragraph(para.text)
        except Exception as e:
            logging.error(f"Error adding paragraph {para_index+1}: {e}")
            record_error("Splitting", f"Paragraph {para_index+1}", para.text, str(e))
            continue

        # Count page breaks in this paragraph
        try:
            page_breaks = count_page_breaks(para)
        except Exception as e:
            logging.error(f"Error counting page breaks in paragraph {para_index+1}: {e}")
            record_error("Splitting", f"Paragraph {para_index+1}", para.text, str(e))
            page_breaks = 0

        if page_breaks > 0:
            global_page += page_breaks

        # If we've reached the segment page limit, save this segment
        if global_page >= segment_start_page + segment_page_count:
            segment_filename = os.path.join(output_dir, f"segment_{segment_index:03d}.docx")
            try:
                current_segment_doc.save(segment_filename)
                logging.info(f"Saved splitted segment: {segment_filename}")
                splitted_files.append(segment_filename)
            except Exception as e:
                logging.error(f"Error saving segment {segment_index}: {e}")
                record_error("Splitting", f"Segment {segment_index}", "", str(e))

            segment_index += 1
            segment_start_page = global_page
            current_segment_doc = Document()

    # Save any remaining paragraphs as the final segment
    if len(current_segment_doc.paragraphs) > 0:
        segment_filename = os.path.join(output_dir, f"segment_{segment_index:03d}.docx")
        try:
            current_segment_doc.save(segment_filename)
            logging.info(f"Saved final splitted segment: {segment_filename}")
            splitted_files.append(segment_filename)
        except Exception as e:
            logging.error(f"Error saving final segment {segment_index}: {e}")
            record_error("Splitting", f"Final Segment {segment_index}", "", str(e))

    return splitted_files


# --------------------------------------------------
# Main
# --------------------------------------------------
def main():
    """
    Usage: python3 split_docx.py <input_docx> <output_folder> <segment_page_count>
    Example: python3 split_docx.py sample.docx splitted 5
    """
    if len(sys.argv) != 4:
        print("Usage: python3 split_docx.py <input_docx> <output_folder> <segment_page_count>")
        sys.exit(1)

    input_docx = sys.argv[1]
    output_folder = sys.argv[2]
    try:
        segment_page_count = int(sys.argv[3])
    except ValueError:
        print("Error: segment_page_count must be an integer.")
        sys.exit(1)

    if not os.path.isfile(input_docx):
        print(f"Error: Input file not found => {input_docx}")
        sys.exit(1)

    splitted_files = split_docx(input_docx, segment_page_count, output_folder)
    print(f"Splitting completed. {len(splitted_files)} file(s) created in '{output_folder}'.")


if __name__ == "__main__":
    main()
