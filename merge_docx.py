#!/usr/bin/env python3
import os
import sys
import glob
import re
from docx import Document

def get_page_number(file_path):
    """
    Extract the page number from filenames like:
    page-36_translated_1-5_20250320_191556.docx
    """
    base = os.path.basename(file_path)
    match = re.search(r'page-(\d+)_translated', base)
    return int(match.group(1)) if match else 0

def merge_docx(files, output_path):
    merged_doc = Document()
    # Remove the default empty paragraph if it exists.
    if merged_doc.paragraphs and not merged_doc.paragraphs[0].text.strip():
        p = merged_doc.paragraphs[0]._element
        p.getparent().remove(p)
    
    for idx, file in enumerate(files):
        print(f"Merging {file} ...")
        sub_doc = Document(file)
        # Append each element from the sub document.
        for element in sub_doc.element.body:
            merged_doc.element.body.append(element)
        # Add a page break between documents (except after the last one)
        if idx < len(files) - 1:
            merged_doc.add_page_break()
    
    merged_doc.save(output_path)
    print(f"\nMerged document saved as: {output_path}")

def main():
    if len(sys.argv) != 2:
        print("Usage: python3 merge_docx.py <directory_with_docx_files>")
        sys.exit(1)
    
    directory = sys.argv[1]
    # Only include DOCX files that contain '_translated_' in the filename.
    pattern = os.path.join(directory, "*_translated_*.docx")
    docx_files = glob.glob(pattern)
    
    if not docx_files:
        print(f"No translated DOCX files found in directory: {directory}")
        sys.exit(1)
    
    # Sort files by the page number extracted from the filename.
    docx_files = sorted(docx_files, key=get_page_number)
    
    output_file = "merged.docx"
    merge_docx(docx_files, output_file)

if __name__ == "__main__":
    main()
