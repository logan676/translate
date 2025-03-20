#!/usr/bin/env python3
from docx import Document
from docx.oxml.ns import qn

def contains_page_break(paragraph):
    # Check each run in the paragraph for a page break element
    for run in paragraph.runs:
        # Look for the <w:br w:type="page"/> element in the run's XML
        if run._element.find('.//w:br[@w:type="page"]') is not None:
            return True
    return False

# Load the original document
doc = Document('my_document.docx')
new_doc = Document()
page_number = 1

for para in doc.paragraphs:
    new_para = new_doc.add_paragraph()
    # Copy runs to preserve formatting (simplified version)
    for run in para.runs:
        new_run = new_para.add_run(run.text)
        # (You can also copy style properties if needed)
    # Check for a page break in this paragraph
    if contains_page_break(para):
        # Save the current document as one page
        new_doc.save(f'page_{page_number}.docx')
        page_number += 1
        # Start a new document for the next page
        new_doc = Document()

# Save any remaining content (if last page did not end with a break)
if new_doc.paragraphs:
    new_doc.save(f'page_{page_number}.docx')
