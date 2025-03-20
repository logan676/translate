#!/usr/bin/env python3

import sys
import os
import datetime
import time
import copy
from docx.oxml import OxmlElement
from docx import Document
from tqdm import tqdm
from tenacity import retry, wait_exponential, stop_after_attempt
import json
import re
from docx.enum.text import WD_BREAK
import traceback
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

# For Google Translate
from googletrans import Translator

TRANSLATE_FONT_SIZE = Pt(9)  # 9号字
TRANSLATE_ITALIC = True      # 斜体
SAVE_INTERVAL = 5            # 每处理5个段落保存一次
PROGRESS_FILE = "translation_progress.json"
PAGE_BREAK_INTERVAL = 3      # 每3段落插入分页符

PAGES_PER_SEGMENT = 5        # How many pages to include in each output before moving on

def run_has_page_break(run):
    """
    Checks if the run contains a page break by examining its XML.
    """
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for br in run._element.findall('.//w:br', namespaces=nsmap):
        if br.get(qn('w:type')) == 'page':
            return True
    return False

def assign_page_numbers(doc):
    """
    Assign a page number to each paragraph by detecting page breaks in runs.
    Returns a list of tuples: [(page_number, paragraph), ...] plus the total page count.
    """
    page_num = 1
    paragraphs_with_page = []

    for paragraph in doc.paragraphs:
        paragraphs_with_page.append((page_num, paragraph))
        for run in paragraph.runs:
            if run_has_page_break(run):
                page_num += 1
                break

    return paragraphs_with_page, page_num

class RealtimeDocWriter:
    """
    Manages incremental writing to a temporary Document,
    along with saving/resuming progress.
    """
    def __init__(self, src_path):
        self.src_doc = Document(src_path)
        self.temp_doc = Document()
        self.progress = self._load_progress()

    def add_paragraph(self, style_clone_from=None, paragraph_text=""):
        """
        添加段落并返回段落对象，以便设置 run 或其他属性。
        style_clone_from: 用来克隆段落格式的源段落
        paragraph_text:   初始写入的新段落文本 
        """
        new_para = self.temp_doc.add_paragraph()

        # 如果需要克隆样式，就调用
        if style_clone_from is not None:
            clone_paragraph_style(style_clone_from, new_para)

        # 如果需要写初始文本
        if paragraph_text:
            new_para.add_run(paragraph_text)
        return new_para
    
    def add_page_break(self):
        self.temp_doc.add_page_break()

    def save(self, output_path):
        self.temp_doc.save(output_path)

    def _load_progress(self):
        try:
            with open(PROGRESS_FILE, 'r') as f:
                return json.load(f)
        except:
            return {'processed_para': 0, 'processed_tables': 0}

    def save_progress(self):
        with open(PROGRESS_FILE, 'w') as f:
            json.dump(self.progress, f)


def should_add_pagebreak(para, count):
    """智能分页判断（示例逻辑，暂未使用）"""
    if para.style.name.startswith('Heading'):
        return True
    return count % 5 == 0

def group_into_segments(paragraphs_with_page):
    """
    Group paragraphs into segments, each containing up to PAGES_PER_SEGMENT pages.
    Returns a dict: {segment_index: [(page_num, paragraph), ...]}.
    """
    segments = {}
    for page_num, paragraph in paragraphs_with_page:
        segment_index = ((page_num - 1) // PAGES_PER_SEGMENT) + 1
        if segment_index not in segments:
            segments[segment_index] = []
        segments[segment_index].append((page_num, paragraph))

    return segments

def init_translator():
    """
    Initialize and return a Googletrans Translator instance.
    """
    translator = Translator()
    return translator

def clone_paragraph_style(src_para, dest_para):
    """
    Safely clone paragraph formatting from src_para to dest_para.
    """
    src_format = src_para.paragraph_format
    dest_format = dest_para.paragraph_format
    
    dest_format.alignment = src_format.alignment
    dest_format.left_indent = src_format.left_indent
    dest_format.right_indent = src_format.right_indent
    dest_format.space_before = src_format.space_before
    dest_format.space_after = src_format.space_after
    dest_format.line_spacing = src_format.line_spacing
    dest_format.widow_control = src_format.widow_control
    
    # Clone tab stops
    try:
        dest_tabs = dest_para.paragraph_format.tab_stops
        while len(dest_tabs) > 0:
            dest_tabs[0].delete()
        
        src_tabs = src_para.paragraph_format.tab_stops
        for src_tab in src_tabs:
            dest_tabs.add_tab_stop(
                src_tab.position,
                src_tab.alignment,
                src_tab.leader
            )
    except Exception as e:
        print(f"制表符克隆失败: {str(e)}")

@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=2, max=10))
def translate_text_with_retry(translator, text):
    """
    Wraps translate_text with tenacity-based retry on failure.
    """
    return translate_text(translator, text)

def translate_text(translator, text, progress_info=""):
    """
    Translates Chinese text into English using Googletrans.
    """
    if not text.strip():
        return ""  # Skip empty text
    
    start_time = time.time()
    try:
        # Translate from Chinese to English
        result = translator.translate(text, src='zh-cn', dest='en')
        translated_text = result.text.strip()
    except Exception as e:
        translated_text = f"[Translation Error: {str(e)}]"
    end_time = time.time()
    time_cost = end_time - start_time

    # Print debug log to console (not included in final output text)
    print(
        "\n--- Debug Log for Translation Request ---\n"
        f"Original Text: {text}\n"
        f"Translated   : {translated_text}\n"
        f"Time Cost    : {time_cost:.2f} seconds\n"
        f"Progress     : {progress_info}\n"
        "--- End Debug Log ---\n"
    )

    return translated_text

def process_tables(doc, translator, writer, progress):
    """
    Demonstration stub for processing tables. 
    Adjust or expand based on your actual logic for table translation.
    """
    tables = doc.tables if hasattr(doc, 'tables') else []
    table_iter = tqdm(
        tables, 
        desc="翻译表格", 
        initial=progress.get('tables', 0),
        position=0,
        leave=True
    )
    
    for table_idx, table in enumerate(table_iter, start=1):
        if table_idx <= progress.get('tables', 0):
            continue
        try:
            # Example logic: For each row, translate each cell
            for row in table.rows:
                for cell in row.cells:
                    original_text = cell.text
                    translated = translate_text_with_retry(translator, original_text)
                    cell.text = translated

                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = TRANSLATE_FONT_SIZE
                            run.italic = TRANSLATE_ITALIC

            # Optionally add a page break after every 2 tables, etc.
            if table_idx % 2 == 0:
                writer.add_page_break()
            
            # Incremental save
            if table_idx % SAVE_INTERVAL == 0:
                writer.save(".temp_output.docx")
                progress['tables'] = table_idx
                writer.save_progress()

        except Exception as e:
            print(f"表格{table_idx}处理失败: {str(e)}")
            traceback.print_exc()
            continue

def process_paragraphs(doc, translator, writer, progress):
    """
    处理段落并将翻译结果增量写入文档。
    """
    para_iter = tqdm(
        doc.paragraphs,
        desc="翻译段落",
        initial=progress.get('paragraph', 0),
        position=1,
        bar_format="\033[92m{l_bar}{bar}\033[0m| {n_fmt}/{total_fmt}",
        leave=False
    )

    for para_idx, src_para in enumerate(para_iter, start=1):
        if para_idx <= progress.get('paragraph', 0):
            continue

        try:
            dest_para = writer.add_paragraph(style_clone_from=src_para)

            if src_para.text.strip():
                original_run = dest_para.add_run(src_para.text)
            
            translated = translate_text_with_retry(translator, src_para.text)
            trans_run = dest_para.add_run('\n' + translated)
            trans_run.font.size = TRANSLATE_FONT_SIZE
            trans_run.italic = TRANSLATE_ITALIC
            trans_run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

            if (para_idx % PAGE_BREAK_INTERVAL == 0) or src_para.style.name.startswith('Heading'):
                writer.add_page_break()

            if para_idx % SAVE_INTERVAL == 0:
                writer.save(".temp_output.docx")
                progress['paragraph'] = para_idx
                writer.save_progress()

        except Exception as e:
            print(f"\n段落 {para_idx} 处理失败: {str(e)}")
            traceback.print_exc()
            continue

def process_segment(translator, segment_paragraphs, segment_range, input_path):
    """
    (Optional) Demonstration of translating paragraphs in segments to TXT.
    This function is not fully integrated if you only want DOCX output.
    """
    print(f"  >> Processing Pages {segment_range[0]}–{segment_range[1]} "
          f"({len(segment_paragraphs)} paragraphs in this segment)")

    base, ext = os.path.splitext(input_path)
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = f"{base}_translated_{segment_range[0]}-{segment_range[1]}_{timestamp}.txt"

    try:
        with open(output_path, "w", encoding="utf-8") as out_file:
            for idx, (page_num, paragraph) in enumerate(segment_paragraphs, start=1):
                raw_text = re.sub(r'\s+', ' ', paragraph.text).strip()
                if len(raw_text) < 1:
                    print(f"跳过空白段落: 第{idx}段")
                    continue    
                print(f"    - Paragraph {idx}/{len(segment_paragraphs)}, Page {page_num}")

                lines = raw_text.split('\n')
                for line_idx, line in enumerate(lines, start=1):
                    line_text = line.strip()
                    if not line_text:
                        continue

                    progress_info = (f"Paragraph {idx}/{len(segment_paragraphs)}, "
                                     f"Page {page_num}, Line {line_idx}/{len(lines)}")
                    print(f"       * Translating line {line_idx} of {len(lines)} in paragraph {idx}")

                    translated_line = translate_text(translator, line_text, progress_info)

                    out_file.write(line_text + "\n")
                    out_file.write(translated_line + "\n\n")
                    out_file.flush()

        print(f"  >> Segment saved to: {output_path}\n")

    except Exception as e:
        print(f"  !! Error saving segment {segment_range[0]}–{segment_range[1]}: {e}")

def main():
    """
    Script workflow:
      1. Read input DOCX file from argv.
      2. Optionally process tables first.
      3. Process paragraphs with incremental save.
      4. Save the final docx.
    """
    if len(sys.argv) != 2:
        print("Usage: python3 translate_docx.py <path_to_docx_file>")
        sys.exit(1)

    input_path = sys.argv[1]

    if not os.path.isfile(input_path):
        print(f"Error: File not found => {input_path}")
        sys.exit(1)

    # Check for existing progress
    if os.path.exists(PROGRESS_FILE):
        with open(PROGRESS_FILE, 'r') as f:
            progress = json.load(f)
        choice = input(
            f"发现未完成的任务（已处理段落 {progress.get('paragraph', 0)}/"
            f"表格 {progress.get('tables', 0)}），是否继续？[y/N] "
        ).lower()
        if choice != 'y':
            os.remove(PROGRESS_FILE)
            progress = {'paragraph': 0, 'tables': 0}
    else:
        progress = {'paragraph': 0, 'tables': 0}

    try:
        doc = Document(input_path)
        translator = init_translator()
        writer = RealtimeDocWriter(input_path)

        # Process tables first (if needed)
        process_tables(doc, translator, writer, progress)
        # Then process paragraphs
        process_paragraphs(doc, translator, writer, progress)

        # Final save
        output_path = os.path.splitext(input_path)[0] + "_translated.docx"
        writer.save(output_path)
        print(f"\n成功生成翻译文档: {output_path}")

    except Exception as e:
        print(f"\n处理中断: {str(e)}")
        print(f"当前进度已保存至 {PROGRESS_FILE}")
    finally:
        # Optionally remove the progress file if you want to reset after each run
        # If you prefer to keep it until the next run, comment this out.
        if os.path.exists(PROGRESS_FILE):
            os.remove(PROGRESS_FILE)

if __name__ == "__main__":
    main()
