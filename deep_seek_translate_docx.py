#!/usr/bin/env python3

import sys
import os
import datetime
import time
import copy
from docx.oxml import OxmlElement
from docx import Document
from tqdm import tqdm
from tenacity import retry, wait_exponential, stop_after_attempt
import json
import re
from docx.enum.text import WD_BREAK
import traceback
from docx import Document
from docx.shared import Pt, RGBColor  # 确保同时导入Pt和RGBColor
from docx.oxml.ns import qn

TRANSLATE_FONT_SIZE = Pt(9)  # 9号字
TRANSLATE_ITALIC = True      # 斜体
SAVE_INTERVAL = 5  # 每处理5个段落保存一次
PROGRESS_FILE = "translation_progress.json"
PAGE_BREAK_INTERVAL = 3      # 每3段落插入分页符

# If you normally use the official OpenAI library, you can do:
#   import openai
#   openai.api_base = DEESEEK_BASE_URL
#   openai.api_key = DEESEEK_API_KEY
import openai

# ---------------------------------------------------------
# DeepSeek / OpenAI Configuration
# ---------------------------------------------------------
DEESEEK_API_KEY = "sk-8df2d0cbcd594a349762d33de5b9df3f"
DEESEEK_BASE_URL = "https://api.deepseek.com"
MODEL_NAME = "deepseek-reasoner"

# How many pages to include in each output before moving on
PAGES_PER_SEGMENT = 5
# ---------------------------------------------------------

def init_client():
    """
    Initialize and return the OpenAI (DeepSeek) client.
    """
    openai.api_key = DEESEEK_API_KEY
    openai.api_base = DEESEEK_BASE_URL
    return openai

def run_has_page_break(run):
    """
    Checks if the run contains a page break by examining its XML.
    """
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for br in run._element.findall('.//w:br', namespaces=nsmap):
        if br.get(qn('w:type')) == 'page':
            return True
    return False

def assign_page_numbers(doc):
    """
    Assign a page number to each paragraph by detecting page breaks in runs.
    Returns a list of tuples: [(page_number, paragraph), ...] plus the total page count.
    """
    page_num = 1
    paragraphs_with_page = []

    for paragraph in doc.paragraphs:
        paragraphs_with_page.append((page_num, paragraph))
        for run in paragraph.runs:
            if run_has_page_break(run):
                page_num += 1
                break

    return paragraphs_with_page, page_num

# +++ 增量写入与断点续传支持 +++
class RealtimeDocWriter:
    def __init__(self, src_path):
        self.src_doc = Document(src_path)
        self.temp_doc = Document()
        self.progress = self._load_progress()
        

    def add_paragraph(self, style_clone_from=None, paragraph_text=""):
        """
        添加段落并返回段落对象，以便设置 run 或其他属性。
        style_clone_from: 用来克隆段落格式的源段落
        paragraph_text:   初始写入的新段落文本 
        """
        
        new_para = self.temp_doc.add_paragraph()

        # 如果需要克隆样式，就调用
        if style_clone_from is not None:
            clone_paragraph_style(style_clone_from, new_para)
        
        # 如果需要写初始文本
        if paragraph_text:
            new_para.add_run(paragraph_text)
        return new_para
    
    def add_page_break(self):
        self.temp_doc.add_page_break()


    def save(self, output_path):
        self.temp_doc.save(output_path)

    def _load_progress(self):
        try:
            with open(PROGRESS_FILE, 'r') as f:
                return json.load(f)
        except:
            return {'processed_para': 0, 'processed_tables': 0}
    
    def save_progress(self):
        with open(PROGRESS_FILE, 'w') as f:
            json.dump(self.progress, f)

# +++ 分页控制逻辑 +++
def should_add_pagebreak(para, count):
    """智能分页判断（网页3/5逻辑）"""
    if para.style.name.startswith('Heading'):
        return True
    return count % 5 == 0

def group_into_segments(paragraphs_with_page):
    """
    Group paragraphs into segments, each containing up to PAGES_PER_SEGMENT pages.
    Returns a dict: {segment_index: [(page_num, paragraph), ...]}.
    """
    segments = {}
    for page_num, paragraph in paragraphs_with_page:
        segment_index = ((page_num - 1) // PAGES_PER_SEGMENT) + 1
        if segment_index not in segments:
            segments[segment_index] = []
        segments[segment_index].append((page_num, paragraph))

    return segments

def translate_text(client, text, progress_info=""):
    """
    Translates the given Chinese text into English using DeepSeek / OpenAI.
    The system message enforces domain context and instructs the model to return
    only the translated text with no extra explanations or summaries.
    """
    system_message = (
        "你是一位高級譯者，專門負責機電系統領域的中英翻譯。"
        "以下是機電系統工程的行業背景與專業名詞依據：\n"
        "（1）機：空調、給水、排水（含雨排水、污廢排水）、消防、防火填塞；\n"
        "（2）電：電力（含接地與避雷系統等）、弱電、通訊、安全、能源（柴油發電機組的燃料來源）。\n\n"
        "請僅提供譯文，不要提供任何解釋、提示或總結。"
    )

    start_time = time.time()
    try:
        # The user message is simply the text to be translated,
        # with no extra instructions or disclaimers.
        response = client.ChatCompletion.create(
            model=MODEL_NAME,
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": text}
            ],
            stream=False
        )
        translated_text = response.choices[0].message.content.strip()
    except Exception as e:
        translated_text = f"[Translation Error: {str(e)}]"
    end_time = time.time()
    time_cost = end_time - start_time

    # Print debug log to console (not included in the final output text).
    print(
        "\n--- Debug Log for Translation Request ---\n"
        f"Original Text: {text}\n"
        f"Translated  : {translated_text}\n"
        f"Time Cost   : {time_cost:.2f} seconds\n"
        f"Progress    : {progress_info}\n"
        "--- End Debug Log ---\n"
    )

    return translated_text

def clone_paragraph_style(src_para, dest_para):
    """安全克隆段落样式（解决paragraph_format无setter问题）"""
    # 克隆段落格式属性（逐个属性复制）
    src_format = src_para.paragraph_format
    dest_format = dest_para.paragraph_format
    
    # 基础格式
    dest_format.alignment = src_format.alignment
    dest_format.left_indent = src_format.left_indent
    dest_format.right_indent = src_format.right_indent
    dest_format.space_before = src_format.space_before
    dest_format.space_after = src_format.space_after
    dest_format.line_spacing = src_format.line_spacing
    dest_format.widow_control = src_format.widow_control
    
    # 修正制表符克隆逻辑
    try:
        # 清空目标制表符（正确方法：逐个删除）
        dest_tabs = dest_para.paragraph_format.tab_stops
        while len(dest_tabs) > 0:
            dest_tabs[0].delete()  # 删除第一个制表符，直到清空
        
        # 克隆源制表符（精确复制）
        src_tabs = src_para.paragraph_format.tab_stops
        for src_tab in src_tabs:
            dest_tabs.add_tab_stop(
                src_tab.position, 
                src_tab.alignment, 
                src_tab.leader
            )
    except Exception as e:
        print(f"制表符克隆失败: {str(e)}")
    

def process_tables(doc, client, writer, progress):
    """处理表格（集成实时样式克隆/分页控制）"""
    from tqdm import tqdm
    
     # 确保tables存在
    tables = doc.tables if hasattr(doc, 'tables') else []

    # 进度条初始化（新增）
    table_iter = tqdm(
        doc.tables, 
        desc="翻译表格", 
        initial=progress['table'],
        position=0,
        leave=True
    )
    
    for table_idx, table in enumerate(table_iter, start=1):
        if table_idx <= progress['table']:
            continue
            
        try:
            # 原始表格样式克隆（增强实现）
            new_row = table.add_row()
            for orig_cell in table.rows[-2].cells:
                new_cell = new_row.cells.add_cell(orig_cell._element)
                
                # 实时翻译（新增重试机制）
                translated = translate_text_with_retry(client, orig_cell.text)
                new_cell.text = translated
                
                # 格式设置（保留原始实现）
                for paragraph in new_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = TRANSLATE_FONT_SIZE
                        run.italic = TRANSLATE_ITALIC
            
            # 分页控制（每2个表格分页）
            if table_idx % 2 == 0:
                writer.add_page_break()
            
            # 增量保存（新增）
            if table_idx % SAVE_INTERVAL == 0:
                writer.save(".temp_output.docx")
                progress['table'] = table_idx
                with open(PROGRESS_FILE, 'w') as f:
                    json.dump(progress, f)
                    
        except Exception as e:
            print(f"表格{table_idx}处理失败: {str(e)}")
            continue

def process_paragraphs(doc, client, writer, progress):
    """
    处理段落并将翻译结果增量写入文档（由 RealtimeDocWriter 管理）。
    - doc: 源文档对象 (Document)
    - client: 已初始化的 openai (DeepSeek) 客户端
    - writer: RealtimeDocWriter 实例，用于实时写入翻译结果
    - progress: dict, 用于记录已处理进度 {"paragraph": 已处理段落数, "table": 已处理表格数}
    """
    from tqdm import tqdm
    
    # 进度条配置
    para_iter = tqdm(
        doc.paragraphs,
        desc="翻译段落",
        initial=progress.get('paragraph', 0),
        position=1,
        bar_format="\033[92m{l_bar}{bar}\033[0m| {n_fmt}/{total_fmt}",
        leave=False
    )
    
    # 遍历文档中的每一个段落
    for para_idx, src_para in enumerate(para_iter, start=1):
        # 如果本段落已在上次进度中被处理过，则跳过
        if para_idx <= progress.get('paragraph', 0):
            continue

        try:
            # 调试：查看原始文本是否为空或含零宽字符
            # print(f"原文: {repr(src_para.text)}")

            # 新建段落（由 RealtimeDocWriter 管理）
            # 并克隆 src_para 的格式到新段落
            dest_para = writer.add_paragraph(style_clone_from=src_para)

            # 先写入原文（如需要在同一个段落里先放原文再放译文）
            if src_para.text.strip():
                original_run = dest_para.add_run(src_para.text)
                # 如果需要也可设置原文字体格式，比如加粗、颜色等
                # original_run.font.bold = True

            # 翻译文本
            translated = translate_text_with_retry(client, src_para.text)

            # 将翻译结果添加到新段落
            trans_run = dest_para.add_run('\n' + translated)  # 换行后再写译文
            trans_run.font.size = TRANSLATE_FONT_SIZE
            trans_run.italic = TRANSLATE_ITALIC
            # 字体颜色示例：RGB(0x42, 0x24, 0xE9)
            trans_run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

            # 分页逻辑示例：每过 PAGE_BREAK_INTERVAL 段，或碰到 Heading，就插入分页符
            if (para_idx % PAGE_BREAK_INTERVAL == 0) or src_para.style.name.startswith('Heading'):
                writer.add_page_break()

            # 每处理一定段落，就增量保存一次
            if para_idx % SAVE_INTERVAL == 0:
                writer.save(".temp_output.docx")
                progress['paragraph'] = para_idx
                writer.save_progress()

        except Exception as e:
            print(f"\n段落 {para_idx} 处理失败: {str(e)}")
            traceback.print_exc()
            continue

# 新增重试装饰器（添加到translate_text函数上方）
@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=2, max=10))
def translate_text_with_retry(client, text):
    return translate_text(client, text)

def process_segment(client, segment_paragraphs, segment_range, input_path):
    """
    Translates paragraphs in a segment line by line, and outputs results to a TXT file.
    """
    print(f"  >> Processing Pages {segment_range[0]}–{segment_range[1]} "
          f"({len(segment_paragraphs)} paragraphs in this segment)")

    base, ext = os.path.splitext(input_path)
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = f"{base}_translated_{segment_range[0]}-{segment_range[1]}_{timestamp}.txt"

    try:
        with open(output_path, "w", encoding="utf-8") as out_file:
            for idx, (page_num, paragraph) in enumerate(segment_paragraphs, start=1):
                raw_text = re.sub(r'\s+', ' ', paragraph.text).strip()  # 替换所有空白字符为单空格
                if len(raw_text) < 1:
                    print(f"跳过空白段落: 第{idx}段")
                    continue    
                print(f"    - Paragraph {idx}/{len(segment_paragraphs)}, Page {page_num}")

                lines = raw_text.split('\n')
                for line_idx, line in enumerate(lines, start=1):
                    line_text = line.strip()
                    if not line_text:
                        continue

                    # Indicate progress in console
                    progress_info = (f"Paragraph {idx}/{len(segment_paragraphs)}, "
                                     f"Page {page_num}, Line {line_idx}/{len(lines)}")
                    print(f"       * Translating line {line_idx} of {len(lines)} in paragraph {idx}")

                    # Perform translation
                    translated_line = translate_text(client, line_text, progress_info)

                    # Write the translated line to the TXT file
                    out_file.write(line_text + "\n")
                    out_file.write(translated_line + "\n")
                    out_file.write("\n")  # extra newline for separation
                    out_file.flush()

        print(f"  >> Segment saved to: {output_path}\n")

    except Exception as e:
        print(f"  !! Error saving segment {segment_range[0]}–{segment_range[1]}: {e}")

def main():
    """
    Script workflow:
      1. Read input DOCX file from argv.
      2. Assign page numbers to paragraphs using page breaks.
      3. Group paragraphs into segments of PAGES_PER_SEGMENT pages.
      4. For each segment, translate line-by-line and save a partial TXT file.
    """
    if len(sys.argv) != 2:
        print("Usage: python3 translate_docx.py <path_to_docx_file>")
        sys.exit(1)

    input_path = sys.argv[1]

     # 断点续传检测（新增）
    if os.path.exists(PROGRESS_FILE):
        with open(PROGRESS_FILE, 'r') as f:
            progress = json.load(f)
        choice = input(f"发现未完成的任务（已处理段落{progress['paragraph']}/表格{progress['table']}），是否继续？[y/N] ").lower()
        if choice != 'y':
            os.remove(PROGRESS_FILE)
            progress = {'paragraph': 0, 'table': 0}
    else:
        progress = {'paragraph': 0, 'table': 0}


    if not os.path.isfile(input_path):
        print(f"Error: File not found => {input_path}")
        sys.exit(1)

    try:
        # 初始化文档（新增延迟加载优化）
        doc = Document(input_path)
        client = init_client()
        
        # 创建带进度条的写入器（新增）
        writer = RealtimeDocWriter(input_path)
        
        # 处理顺序优化（先表格后段落）
        process_tables(doc, client, writer, progress)
        process_paragraphs(doc, client, writer, progress)
        
        # 最终保存（新增异常保护）
        output_path = os.path.splitext(input_path)[0] + "_translated.docx"
        writer.save(output_path)
        print(f"\成功生成翻译文档：: {output_path}")
        
    except Exception as e:
        print(f"\n处理中断: {str(e)}")
        print(f"当前进度已保存至 {PROGRESS_FILE}")
    finally:
        if os.path.exists(PROGRESS_FILE):
            os.remove(PROGRESS_FILE)
if __name__ == "__main__":
    main()
